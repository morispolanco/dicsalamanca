import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Diccionario Económico de la Escuela de Salamanca", page_icon="📚", layout="wide")

# Función para crear la columna de información
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Económico basado en el pensamiento de la Escuela de Salamanca. Permite a los usuarios obtener definiciones de términos económicos según la interpretación de diversos autores de esta escuela.

    ### Cómo usar la aplicación:

    1. Elija un término económico de la lista predefinida o proponga su propio término.
    2. Seleccione uno o más autores de la Escuela de Salamanca.
    3. Haga clic en "Obtener definición" para generar las definiciones.
    4. Lea las definiciones y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 25 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Económico de la Escuela de Salamanca* [Aplicación web]. https://ecsalamanca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Título de la aplicación
st.title("Diccionario Económico de la Escuela de Salamanca")

# Crear un diseño de dos columnas
col1, col2 = st.columns([1, 2])

# Columna de información
with col1:
    crear_columna_info()

# Columna principal
with col2:
    # Acceder a las claves de API de los secretos de Streamlit
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # Lista de términos económicos
    terminos_economicos = [
        "Ahorro", "Ahorro y préstamo", "Arbitraje", "Beneficio", "Bien común", "Bienestar", "Bancarrota", "Capital",
        "Comercio justo", "Competencia", "Confianza", "Consumo", "Contrato", "Corruptibilidad", "Coste de oportunidad",
        "Costes", "Crédito", "Crédito bancario", "Deuda", "Déficit", "Demandas básicas", "Descuento", 
        "Desigualdad económica", "Desvalorización", "División del trabajo", "Dinero fiduciario", "Dinero metálico",
        "Dignidad humana", "Derecho natural", "Derechos de propiedad", "Derechos humanos", "Economía moral",
        "Emancipación económica", "Emprendimiento", "Equidad", "Equilibrio de mercado", "Especulación", "Excedente",
        "Explotación", "Exportación", "Fiscalidad", "Fraude", "Función empresarial", "Génesis del capital",
        "Importación", "Innovación", "Intercambio", "Intercambio internacional", "Intercambio voluntario",
        "Inversión", "Interés", "Intervención estatal", "Justicia distributiva", "Justicia conmutativa",
        "Justicia social", "Ley de la oferta y la demanda", "Libertad económica", "Libertad de contratación",
        "Mercado", "Medio de cambio", "Moneda", "Monopolio", "Movilidad social", "Orden natural", "Paridad monetaria",
        "Poder adquisitivo", "Precio justo", "Precios de mercado", "Producción", "Productividad", 
        "Propiedad comunitaria", "Propiedad privada", "Prosperidad económica", "Prosperidad sostenible",
        "Prudencia económica", "Reciprocidad", "Regulación", "Rentabilidad", "Reserva de valor",
        "Respuesta del mercado", "Responsabilidad moral", "Responsabilidad personal", "Riesgo", "Salud económica",
        "Solidaridad", "Soberanía", "Subsidiariedad", "Temor de la ley", "Teoría del valor", "Tesorería",
        "Tributación", "Trueque", "Usura", "Valor subjetivo", "Violencia económica"
    ]

    # Lista de autores de la Escuela de Salamanca
    autores_salamanca = [
        "Arias Piñel", "Antonio de Padilla y Meneses", "Bartolomé de Albornoz", "Bartolomé de Medina",
        "Diego de Chaves", "Diego de Covarrubias", "Diego Pérez de Mesa", "Domingo Báñez", "Domingo de Soto",
        "Fernán Pérez de Oliva", "Francisco de Vitoria", "Francisco Sarmiento de Mendoza", "Francisco Suárez",
        "Gregorio de Valencia", "Jerónimo Muñoz", "Juan de Horozco y Covarrubias", "Juan de la Peña",
        "Juan de Matienzo", "Juan de Ribera", "Juan Gil de la Nava", "Leonardus Lessius", "Luis de León",
        "Martín de Azpilcueta", "Martín de Ledesma", "Melchor Cano", "Pedro de Sotomayor", "Tomás de Mercado",
        "Alonso de la Vera Cruz", "Cristóbal de Villalón", "Fernando Vázquez de Menchaca",
        "Francisco Cervantes de Salazar", "Juan de Lugo y Quiroga", "Juan de Salas", "Luis de Molina",
        "Pedro de Aragón", "Pedro de Valencia", "Antonio de Hervías", "Bartolomé de Carranza",
        "Bartolomé de las Casas", "Cristóbal de Fonseca", "Domingo de Salazar", "Domingo de Santo Tomás",
        "Gabriel Vásquez", "Gómez Pereira", "Juan de Mariana", "Juan de Medina", "Juan Pérez de Menacho",
        "Luis de Alcalá", "Luis Saravia de la Calle", "Miguel Bartolomé Salón", "Pedro de Fonseca",
        "Pedro de Oñate", "Rodrigo de Arriaga"
    ]

    def buscar_informacion(query, autor):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {autor} Escuela de Salamanca economía"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.request("POST", url, headers=headers, data=payload)
        return response.json()

    def generar_definicion(termino, autor, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\nAutor: {autor}\n\nProporciona una definición del término económico '{termino}' según el pensamiento de {autor}, un autor de la Escuela de Salamanca. La definición debe ser concisa pero informativa, similar a una entrada de diccionario. Si es posible, incluye una referencia a una obra específica de {autor} que trate este concepto.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.request("POST", url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, definiciones, fuentes):
        doc = Document()
        doc.add_heading('Diccionario Económico - Escuela de Salamanca', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        for autor, definicion in definiciones.items():
            doc.add_heading(f'Definición según {autor}', level=2)
            doc.add_paragraph(definicion)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    # Interfaz de usuario
    st.write("Elige un término económico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_economicos)
    else:
        termino = st.text_input("Ingresa tu propio término económico:")

    # Selección de autores
    st.write("Selecciona uno o más autores de la Escuela de Salamanca (máximo 5):")
    autores_seleccionados = st.multiselect("Autores", autores_salamanca)

    if len(autores_seleccionados) > 5:
        st.warning("Has seleccionado más de 5 autores. Por favor, selecciona un máximo de 5.")
    else:
        if st.button("Obtener definición"):
            if termino and autores_seleccionados:
                with st.spinner("Buscando información y generando definiciones..."):
                    definiciones = {}
                    todas_fuentes = []

                    for autor in autores_seleccionados:
                        # Buscar información relevante
                        resultados_busqueda = buscar_informacion(termino, autor)
                        contexto = "\n".join([result.get('snippet', '') for result in resultados_busqueda.get('organic', [])])

                        # Generar definición
                        definicion = generar_definicion(termino, autor, contexto)
                        definiciones[autor] = definicion

                        # Recopilar fuentes
                        fuentes = [f"{resultado['title']}: {resultado['link']}" for resultado in resultados_busqueda.get('organic', [])[:3]]
                        todas_fuentes.extend(fuentes)

                    # Mostrar definiciones
                    st.write(f"Definiciones de '{termino}':")
                    for autor, definicion in definiciones.items():
                        st.write(f"\nSegún {autor}:")
                        st.write(definicion)

                    # Mostrar fuentes
                    st.write("\nFuentes:")
                    for fuente in todas_fuentes:
                        st.write(f"- {fuente}")

                    # Crear documento DOCX
                    doc = create_docx(termino, definiciones, todas_fuentes)

                    # Guardar el documento DOCX en memoria
                    docx_file = BytesIO()
                    doc.save(docx_file)
                    docx_file.seek(0)

                    # Opción para exportar a DOCX
                    st.download_button(
                        label="Descargar definiciones como DOCX",
                        data=docx_file,
                        file_name=f"definiciones_{termino.lower().replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            elif not termino:
                st.warning("Por favor, selecciona o ingresa un término.")
            elif not autores_seleccionados:
                st.warning("Por favor, selecciona al menos un autor.")
