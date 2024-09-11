import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Asistente de Economía - Escuela de Salamanca", page_icon="📚")

# Título de la aplicación
st.title("Asistente de Economía - Escuela de Salamanca")

# Acceder a las claves de API de los secretos de Streamlit
TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

# Lista de 99 preguntas predefinidas
preguntas_predefinidas = [
    "¿Cuáles son los principales aportes de Francisco de Vitoria a la economía?",
    "¿Qué papel jugó la Escuela de Salamanca en el desarrollo del derecho internacional?",
    "¿Cómo se aborda la teoría del valor en la Escuela de Salamanca?",
    "¿Cuál es la influencia de Martín de Azpilcueta en la teoría monetaria?",
    "¿Cómo se relaciona la teoría del interés con la Escuela de Salamanca?",
    "¿Qué impacto tuvieron los pensadores de la Escuela de Salamanca en la economía moderna?",
    "¿Qué relación existe entre la ética cristiana y la teoría económica de la Escuela de Salamanca?",
    "¿Cuáles son las contribuciones de Domingo de Soto a la teoría económica?",
    "¿Cómo aborda la Escuela de Salamanca el tema de los monopolios?",
    "¿Qué opinaba la Escuela de Salamanca sobre la usura?",
    "¿Cómo influyó la colonización en las teorías económicas de la Escuela de Salamanca?",
    "¿Cuál es la posición de la Escuela de Salamanca respecto a los impuestos y la tributación?",
    "¿Cómo trata la Escuela de Salamanca el concepto de propiedad privada?",
    "¿Qué es el precio justo según los pensadores de la Escuela de Salamanca?",
    "¿Cómo influyó la escolástica en el desarrollo del pensamiento económico?",
    "¿Qué relación hay entre la Escuela de Salamanca y el liberalismo económico?",
    "¿Cómo interpreta la Escuela de Salamanca la inflación monetaria?",
    "¿Cómo afectaron los descubrimientos de oro y plata en América a las teorías económicas de la Escuela de Salamanca?",
    "¿Qué postura tenía la Escuela de Salamanca sobre el comercio internacional?",
    "¿Cómo influyeron los escritos de la Escuela de Salamanca en el desarrollo de la teoría de la banca?",
    "¿Cuál es el legado de Luis de Molina en la teoría del mercado?",
    "¿Cómo aborda la Escuela de Salamanca los contratos y las obligaciones?",
    "¿Cómo se desarrolló el concepto de ley natural en la Escuela de Salamanca?",
    "¿Qué impacto tuvieron los escritos de los escolásticos en la moral económica?",
    "¿Qué papel jugó la justicia distributiva en las teorías de la Escuela de Salamanca?",
    "¿Qué diferencias había entre los pensadores de la Escuela de Salamanca y los escolásticos medievales?",
    "¿Qué aportaciones hizo Juan de Mariana a la economía?",
    "¿Cómo influyeron las ideas de la Escuela de Salamanca en los movimientos de reforma económica del siglo XVI?",
    "¿Cómo abordó la Escuela de Salamanca el tema de la libertad económica?",
    "¿Cuál es la postura de la Escuela de Salamanca sobre la esclavitud?",
    "¿Qué relación había entre el derecho canónico y las ideas económicas de la Escuela de Salamanca?",
    "¿Qué pensaba la Escuela de Salamanca sobre el salario justo?",
    "¿Cuál es la posición de la Escuela de Salamanca sobre el valor de los metales preciosos?",
    "¿Cómo influyó la Escuela de Salamanca en la teoría de la competencia?",
    "¿Cómo aborda la Escuela de Salamanca el concepto de utilidad en la economía?",
    "¿Qué relación hay entre las ideas de la Escuela de Salamanca y el capitalismo?",
    "¿Cómo se desarrolló la teoría del cambio en la Escuela de Salamanca?",
    "¿Cómo se explica la influencia de la teología en las teorías económicas de la Escuela de Salamanca?",
    "¿Cuál fue el papel de Diego de Covarrubias y Leyva en la teoría del valor?",
    "¿Qué relación hay entre las ideas de la Escuela de Salamanca y las reformas económicas posteriores en Europa?",
    "¿Qué aportaciones hizo Tomás de Mercado a las ideas sobre comercio y banca?",
    "¿Cómo trata la Escuela de Salamanca la teoría de los precios de mercado?",
    "¿Cómo influyó la Escuela de Salamanca en la teoría del crédito?",
    "¿Qué impacto tuvo la Escuela de Salamanca en la teoría de los recursos naturales?",
    "¿Cómo se desarrolló el concepto de intercambio voluntario en la Escuela de Salamanca?",
    "¿Qué opinaban los autores de la Escuela de Salamanca sobre el control estatal de la economía?",
    "¿Cómo afectaron los descubrimientos geográficos a las teorías económicas de la Escuela de Salamanca?",
    "¿Qué aportaciones hizo Bartolomé de las Casas a la economía y la justicia social?",
    "¿Qué pensaba la Escuela de Salamanca sobre el derecho a la propiedad comunal?",
    "¿Qué ideas de la Escuela de Salamanca influyeron en el derecho a la autodeterminación de los pueblos indígenas?",
    "¿Cómo influye la moral cristiana en las teorías económicas de la Escuela de Salamanca?",
    "¿Qué relación existe entre la justicia conmutativa y las ideas de la Escuela de Salamanca?",
    "¿Cómo aborda la Escuela de Salamanca el concepto de beneficio en los negocios?",
    "¿Qué rol jugó el debate sobre la legitimidad del poder en las ideas económicas de la Escuela de Salamanca?",
    "¿Qué pensaba la Escuela de Salamanca sobre el sistema de gremios?",
    "¿Qué contribuciones hizo la Escuela de Salamanca a la teoría del contrato social?",
    "¿Cómo se relacionan las ideas de la Escuela de Salamanca con el mercantilismo?",
    "¿Qué opinaba la Escuela de Salamanca sobre la intervención estatal en la economía?",
    "¿Cómo influenció la Escuela de Salamanca a los economistas clásicos como Adam Smith?",
    "¿Qué relación existe entre la ética y la economía en la Escuela de Salamanca?",
    "¿Cómo abordaba la Escuela de Salamanca la relación entre oferta y demanda?",
    "¿Qué ideas de la Escuela de Salamanca anticiparon la teoría moderna del dinero?",
    "¿Cómo influyeron los autores de la Escuela de Salamanca en la teoría del interés compuesto?",
    "¿Qué pensaba la Escuela de Salamanca sobre los tratados comerciales internacionales?",
    "¿Cómo influyó la Escuela de Salamanca en la teoría de los mercados financieros?",
    "¿Cuál es la postura de la Escuela de Salamanca sobre la justicia en el comercio?",
    "¿Cómo abordó la Escuela de Salamanca la relación entre riqueza y bienestar?",
    "¿Cómo trata la Escuela de Salamanca el concepto de ganancia legítima?",
    "¿Cómo se aborda el valor intrínseco de los bienes en la Escuela de Salamanca?",
    "¿Qué influencia tuvo la Escuela de Salamanca en la teoría de la deuda pública?",
    "¿Cuál es la postura de la Escuela de Salamanca sobre el libre mercado?",
    "¿Qué opinaba la Escuela de Salamanca sobre la circulación de moneda?",
    "¿Cómo influyó la Escuela de Salamanca en las políticas económicas coloniales?",
    "¿Qué papel jugó la Escuela de Salamanca en el desarrollo del derecho mercantil?",
    "¿Cómo se abordaba el tema del fraude en las transacciones económicas según la Escuela de Salamanca?",
    "¿Cuál es la influencia de los pensadores de la Escuela de Salamanca en la economía contemporánea?",
    "¿Cómo influyó la Escuela de Salamanca en la ética de los negocios?",
    "¿Qué papel juega la teoría del contrato en las ideas económicas de la Escuela de Salamanca?",
    "¿Qué papel jugó la teología moral en la formulación de las teorías económicas de la Escuela de Salamanca?",
    "¿Qué influencia tuvo la Escuela de Salamanca en la reforma fiscal de la época?",
    "¿Cómo se abordaba el tema de los préstamos en la Escuela de Salamanca?",
    "¿Cuál es el legado de la Escuela de Salamanca en la teoría del comercio internacional?",
    "¿Cómo se abordaba el problema de la inflación en los tiempos de la Escuela de Salamanca?",
    "¿Qué papel jugaba la doctrina del precio justo en la teoría económica de la Escuela de Salamanca?",
    "¿Cuál fue la influencia de los pensadores de la Escuela de Salamanca en la formación de las primeras políticas fiscales?",
    "¿Cómo se relacionan las ideas de la Escuela de Salamanca con la responsabilidad social empresarial?",
    "¿Cómo influenció la Escuela de Salamanca las primeras formas de análisis económico matemático?",
    "¿Cuál fue la contribución de la Escuela de Salamanca al entendimiento de la soberanía económica?",
    "¿Cómo influyó la Escuela de Salamanca en la regulación del comercio de bienes y servicios?",
    "¿Cómo influyó la Escuela de Salamanca en el desarrollo de la teoría del bien común?",
    "¿Qué influencia tuvo la Escuela de Salamanca en las políticas sobre propiedad intelectual en la época?",
    "¿Cuál fue la relación entre la Escuela de Salamanca y la teoría del equilibrio general en la economía?",
    "¿Qué relación existe entre las ideas de la Escuela de Salamanca y las primeras nociones de globalización económica?"
]

def buscar_informacion(query):
    url = "https://google.serper.dev/search"
    payload = json.dumps({
        "q": query + " Escuela de Salamanca economía"
    })
    headers = {
        'X-API-KEY': SERPER_API_KEY,
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    resultados = response.json().get('organic', [])
    
    # Ampliar a los primeros 5 resultados para más contexto
    contexto = "\n".join([result.get('snippet', '') for result in resultados[:5]])
    return contexto, resultados

def generar_respuesta(prompt, contexto):
    url = "https://api.together.xyz/inference"
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": f"Contexto: {contexto}\n\nPregunta: {prompt}\n\nResponde la pregunta de manera extensa, proporcionando un análisis detallado de los conceptos económicos según la Escuela de Salamanca. Incluye referencias a las principales obras y autores. No te limites a un resumen breve, elabora en profundidad la respuesta con al menos 4 fuentes.\n\nRespuesta:",
        "max_tokens": 7000,  # Aumentar el número de tokens
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["Pregunta:"]
    })
    headers = {
        'Authorization': f'Bearer {TOGETHER_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()['output']['choices'][0]['text'].strip()

def create_docx(pregunta, respuesta, fuentes):
    doc = Document()
    doc.add_heading('Asistente de Economía - Escuela de Salamanca', 0)

    doc.add_heading('Pregunta', level=1)
    doc.add_paragraph(pregunta)

    doc.add_heading('Respuesta', level=1)
    doc.add_paragraph(respuesta)

    doc.add_heading('Fuentes', level=1)
    for fuente in fuentes:
        doc.add_paragraph(fuente, style='List Bullet')

    doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

    return doc

# Interfaz de usuario
st.subheader("Selecciona una pregunta o escribe la tuya propia:")

# Menú desplegable para elegir una pregunta
pregunta_seleccionada = st.selectbox("Selecciona una pregunta", preguntas_predefinidas)

# Campo para que el usuario pueda escribir su propia pregunta
pregunta_personalizada = st.text_input("O escribe tu propia pregunta aquí:")

# Usar la pregunta personalizada si se escribe, si no, usar la pregunta seleccionada
pregunta = pregunta_personalizada if pregunta_personalizada else pregunta_seleccionada

if st.button("Obtener respuesta"):
    if pregunta:
        with st.spinner("Buscando información y generando respuesta..."):
            # Buscar información relevante
            contexto, resultados_busqueda = buscar_informacion(pregunta)

            # Generar respuesta
            respuesta = generar_respuesta(pregunta, contexto)

            # Mostrar respuesta
            st.write("Respuesta:")
            st.write(respuesta)

            # Mostrar fuentes
            st.write("Fuentes:")
            fuentes = []
            for resultado in resultados_busqueda[:5]:  # Ampliar a 5 fuentes
                fuente = f"{resultado['title']}: {resultado['link']}"
                st.write(f"- [{resultado['title']}]({resultado['link']})")
                fuentes.append(fuente)

            # Crear documento DOCX
            doc = create_docx(pregunta, respuesta, fuentes)

            # Guardar el documento DOCX en memoria
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)

            # Opción para exportar a DOCX
            st.download_button(
                label="Descargar resultados como DOCX",
                data=docx_file,
                file_name="respuesta_economia_salamanca.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    else:
        st.warning("Por favor, ingresa o selecciona una pregunta.")

# Agregar información en el pie de página
st.markdown("---")
st.markdown("**Nota:** Este asistente utiliza IA para generar respuestas basadas en información disponible en línea sobre la Escuela de Salamanca. "
            "Siempre verifica la información con fuentes académicas para un análisis más profundo de los conceptos económicos.")
