import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Diccionario Económico de la Escuela de Salamanca", page_icon="📚", layout="wide")

# Función para crear la columna de información
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Económico basado en el pensamiento de la Escuela de Salamanca. Permite a los usuarios obtener definiciones de términos económicos o ensayos académicos que incluyen comparaciones con otras escuelas de pensamiento.

    ### Cómo usar la aplicación:

    1. Elija un término económico de la lista predefinida o proponga su propio término.
    2. Seleccione si desea generar un artículo de diccionario o un ensayo académico.
    3. Haga clic en "Generar contenido" para obtener el resultado.
    4. Lea el contenido generado y las fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 25 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Económico de la Escuela de Salamanca* [Aplicación web]. https://ecsalamanca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar contenido basado en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Título de la aplicación
st.title("Diccionario Económico de la Escuela de Salamanca")

# Crear un diseño de dos columnas
col1, col2 = st.columns([1, 2])

# Columna de información
with col1:
    crear_columna_info()

# Columna principal
with col2:
    # Acceder a las claves de API de los secretos de Streamlit
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # Lista de términos económicos
    terminos_economicos = [
        "Ahorro", "Ahorro y préstamo", "Arbitraje", "Beneficio", "Bien común", "Bienestar", "Bancarrota", "Capital",
        "Comercio justo", "Competencia", "Confianza", "Consumo", "Contrato", "Corruptibilidad", "Coste de oportunidad",
        "Costes", "Crédito", "Crédito bancario", "Deuda", "Déficit", "Demandas básicas", "Descuento", 
        "Desigualdad económica", "Desvalorización", "División del trabajo", "Dinero fiduciario", "Dinero metálico",
        "Dignidad humana", "Derecho natural", "Derechos de propiedad", "Derechos humanos", "Economía moral",
        "Emancipación económica", "Emprendimiento", "Equidad", "Equilibrio de mercado", "Especulación", "Excedente",
        "Explotación", "Exportación", "Fiscalidad", "Fraude", "Función empresarial", "Génesis del capital",
        "Importación", "Innovación", "Intercambio", "Intercambio internacional", "Intercambio voluntario",
        "Inversión", "Interés", "Intervención estatal", "Justicia distributiva", "Justicia conmutativa",
        "Justicia social", "Ley de la oferta y la demanda", "Libertad económica", "Libertad de contratación",
        "Mercado", "Medio de cambio", "Moneda", "Monopolio", "Movilidad social", "Orden natural", "Paridad monetaria",
        "Poder adquisitivo", "Precio justo", "Precios de mercado", "Producción", "Productividad", 
        "Propiedad comunitaria", "Propiedad privada", "Prosperidad económica", "Prosperidad sostenible",
        "Prudencia económica", "Reciprocidad", "Regulación", "Rentabilidad", "Reserva de valor",
        "Respuesta del mercado", "Responsabilidad moral", "Responsabilidad personal", "Riesgo", "Salud económica",
        "Solidaridad", "Soberanía", "Subsidiariedad", "Temor de la ley", "Teoría del valor", "Tesorería",
        "Tributación", "Trueque", "Usura", "Valor subjetivo", "Violencia económica"
    ]

    def buscar_informacion(query):
        url = "https://api.serply.io/v1/scholar"
        params = {
            "q": f"{query} Escuela de Salamanca economía"
        }
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.get(url, headers=headers, params=params)
        return response.json()

    def generar_contenido(termino, tipo_contenido):
        url = "https://api.together.xyz/inference"
        if tipo_contenido == "Generar artículo de diccionario":
            prompt = f"""Crea un artículo de diccionario para el término económico '{termino}' basado en el pensamiento de la Escuela de Salamanca. 
            Incluye definiciones y discusiones de varios autores de esta escuela, citando sus obras específicas. 
            El artículo debe ser conciso pero informativo, similar a una entrada de diccionario enciclopédico."""
        else:
            prompt = f"""Escribe un ensayo académico sobre el término económico '{termino}' desde la perspectiva de la Escuela de Salamanca. 
            Incluye una discusión de varios autores de esta escuela, citando sus obras. 
            Además, compara el concepto con la interpretación en la Doctrina Social de la Iglesia y los principios de la Escuela Austríaca de Economía. 
            Proporciona un análisis crítico y comparativo de estas perspectivas."""

        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": prompt,
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, contenido, fuentes, tipo_contenido):
        doc = Document()
        doc.add_heading('Diccionario Económico - Escuela de Salamanca', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading(tipo_contenido, level=1)
        doc.add_paragraph(contenido)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    # Interfaz de usuario
    st.write("Elige un término económico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_economicos)
    else:
        termino = st.text_input("Ingresa tu propio término económico:")

    # Selección del tipo de contenido
    tipo_contenido = st.radio("Selecciona el tipo de contenido a generar:", ["Generar artículo de diccionario", "Generar ensayo académico"])

    if st.button("Generar contenido"):
        if termino:
            with st.spinner("Buscando información y generando contenido..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                
                # Generar contenido
                contenido = generar_contenido(termino, tipo_contenido)

                # Mostrar contenido
                st.write(f"{tipo_contenido} para '{termino}':")
                st.write(contenido)

                # Recopilar y mostrar fuentes
                fuentes = [f"{resultado['title']}: {resultado['link']}" for resultado in resultados_busqueda.get('results', [])[:5]]
                st.write("\nFuentes:")
                for fuente in fuentes:
                    st.write(f"- {fuente}")

                # Crear documento DOCX
                doc = create_docx(termino, contenido, fuentes, tipo_contenido)

                # Guardar el documento DOCX en memoria
                docx_file = BytesIO()
                doc.save(docx_file)
                docx_file.seek(0)

                # Opción para exportar a DOCX
                st.download_button(
                    label="Descargar contenido como DOCX",
                    data=docx_file,
                    file_name=f"{tipo_contenido.lower().replace(' ', '_')}_{termino.lower().replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        else:
            st.warning("Por favor, selecciona o ingresa un término.")
