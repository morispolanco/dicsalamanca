import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from io import BytesIO
import re

# Configuración de la página
st.set_page_config(page_title="Diccionario Económico de la Escuela de Salamanca", page_icon="📚", layout="wide")

# Función para crear la columna de información
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Económico basado en el pensamiento de la Escuela de Salamanca. Permite a los usuarios obtener definiciones de términos económicos o ensayos académicos que incluyen comparaciones con otras escuelas de pensamiento.

    ### Cómo usar la aplicación:

    1. Elija un término económico de la lista predefinida o proponga su propio término.
    2. Seleccione si desea generar un artículo de diccionario o un ensayo académico.
    3. Haga clic en "Generar contenido" para obtener el resultado.
    4. Lea el contenido generado con enlaces a las fuentes citadas.
    5. Si lo desea, descargue un documento DOCX con toda la información y enlaces activos.

    ### Autor y actualización:
    **Moris Polanco**, 25 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Económico de la Escuela de Salamanca* [Aplicación web]. https://ecsalamanca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar contenido basado en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Título de la aplicación
st.title("Diccionario Económico de la Escuela de Salamanca")

# Crear un diseño de dos columnas
col1, col2 = st.columns([1, 2])

# Columna de información
with col1:
    crear_columna_info()

# Columna principal
with col2:
    # Acceder a las claves de API de los secretos de Streamlit
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPAPI_API_KEY = st.secrets["SERPAPI_API_KEY"]

    # Lista de términos económicos
    terminos_economicos = [
        "Ahorro", "Ahorro y préstamo", "Arbitraje", "Beneficio", "Bien común", "Bienestar", "Bancarrota", "Capital",
        "Comercio justo", "Competencia", "Confianza", "Consumo", "Contrato", "Corruptibilidad", "Coste de oportunidad",
        "Costes", "Crédito", "Crédito bancario", "Deuda", "Déficit", "Demandas básicas", "Descuento", 
        "Desigualdad económica", "Desvalorización", "División del trabajo", "Dinero fiduciario", "Dinero metálico",
        "Dignidad humana", "Derecho natural", "Derechos de propiedad", "Derechos humanos", "Economía moral",
        "Emancipación económica", "Emprendimiento", "Equidad", "Equilibrio de mercado", "Especulación", "Excedente",
        "Explotación", "Exportación", "Fiscalidad", "Fraude", "Función empresarial", "Génesis del capital",
        "Importación", "Innovación", "Intercambio", "Intercambio internacional", "Intercambio voluntario",
        "Inversión", "Interés", "Intervención estatal", "Justicia distributiva", "Justicia conmutativa",
        "Justicia social", "Ley de la oferta y la demanda", "Libertad económica", "Libertad de contratación",
        "Mercado", "Medio de cambio", "Moneda", "Monopolio", "Movilidad social", "Orden natural", "Paridad monetaria",
        "Poder adquisitivo", "Precio justo", "Precios de mercado", "Producción", "Productividad", 
        "Propiedad comunitaria", "Propiedad privada", "Prosperidad económica", "Prosperidad sostenible",
        "Prudencia económica", "Reciprocidad", "Regulación", "Rentabilidad", "Reserva de valor",
        "Respuesta del mercado", "Responsabilidad moral", "Responsabilidad personal", "Riesgo", "Salud económica",
        "Solidaridad", "Soberanía", "Subsidiariedad", "Temor de la ley", "Teoría del valor", "Tesorería",
        "Tributación", "Trueque", "Usura", "Valor subjetivo", "Violencia económica"
    ]

    def buscar_informacion(query):
        url = "https://serpapi.com/search?engine=google_scholar_cite"
        params = {
            "q": f"{query} Escuela de Salamanca economía",
            "api_key": SERPAPI_API_KEY
        }
        response = requests.get(url, params=params)
        return response.json()

    def generar_contenido(termino, tipo_contenido, fuentes):
        url = "https://api.together.xyz/inference"
        fuentes_str = "\n".join([f"- {fuente}" for fuente in fuentes])
        if tipo_contenido == "Generar artículo de diccionario":
            prompt = f"""Crea un artículo de diccionario para el término económico '{termino}' basado en el pensamiento de la Escuela de Salamanca. 
            Incluye definiciones y discusiones de varios autores de esta escuela, citando sus obras específicas. 
            El artículo debe ser conciso pero informativo, similar a una entrada de diccionario enciclopédico.
            Utiliza y cita las siguientes fuentes en tu respuesta:
            {fuentes_str}
            Asegúrate de incluir citas en el texto y una lista de referencias al final. 
            Para cada cita en el texto, usa el formato [Autor, Año] y asegúrate de que corresponda con una entrada en la lista de referencias."""
        else:
            prompt = f"""Escribe un ensayo académico sobre el término económico '{termino}' desde la perspectiva de la Escuela de Salamanca. 
            Incluye una discusión de varios autores de esta escuela, citando sus obras. 
            Además, compara el concepto con la interpretación en la Doctrina Social de la Iglesia y los principios de la Escuela Austríaca de Economía. 
            Proporciona un análisis crítico y comparativo de estas perspectivas.
            Utiliza y cita las siguientes fuentes en tu ensayo:
            {fuentes_str}
            Asegúrate de incluir citas en el texto y una lista de referencias al final. 
            Para cada cita en el texto, usa el formato [Autor, Año] y asegúrate de que corresponda con una entrada en la lista de referencias."""

        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": prompt,
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), '0000FF')
        rPr.append(c)
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def create_docx(termino, contenido, tipo_contenido, fuentes):
        doc = Document()
        doc.add_heading('Diccionario Económico - Escuela de Salamanca', 0)
        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)
        doc.add_heading(tipo_contenido, level=1)
        parrafos = contenido.split('\n\n')
        for parrafo in parrafos:
            p = doc.add_paragraph()
            citas = re.findall(r'\[([^\]]+)\]', parrafo)
            partes = re.split(r'\[([^\]]+)\]', parrafo)
            for i, parte in enumerate(partes):
                if i % 2 == 0:
                    p.add_run(parte)
                else:
                    for fuente in fuentes:
                        if parte.lower() in fuente.lower():
                            enlace = fuente.split(': ')[-1]
                            add_hyperlink(p, enlace, f'[{parte}]')
                            break
                    else:
                        p.add_run(f'[{parte}]')
        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')
        return doc

    # Interfaz de usuario
    st.write("Elige un término económico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_economicos)
    else:
        termino = st.text_input("Ingresa tu propio término económico:")

    tipo_contenido = st.radio("Selecciona el tipo de contenido a generar:", ["Generar artículo de diccionario", "Generar ensayo académico"])

    if st.button("Generar contenido"):
        if termino:
            with st.spinner("Buscando información y generando contenido..."):
                resultados_busqueda = buscar_informacion(termino)
                fuentes = [f"{resultado['title']}: {resultado['link']}" for resultado in resultados_busqueda.get('citations', [])[:5]]
                contenido = generar_contenido(termino, tipo_contenido, fuentes)
                st.write(f"{tipo_contenido} para '{termino}':")
                parrafos = contenido.split('\n\n')
                for parrafo in parrafos:
                    citas = re.findall(r'\[([^\]]+)\]', parrafo)
                    partes = re.split(r'\[([^\]]+)\]', parrafo)
                    nuevo_parrafo = ""
                    for i, parte in enumerate(partes):
                        if i % 2 == 0:
                            nuevo_parrafo += parte
                        else:
                            for fuente in fuentes:
                                if parte.lower() in fuente.lower():
                                    enlace = fuente.split(': ')[-1]
                                    nuevo_parrafo += f'[{parte}]({enlace})'
                                    break
                            else:
                                nuevo_parrafo += f'[{parte}]'
                    st.markdown(nuevo_parrafo)
                    st.write("")  
                doc = create_docx(termino, contenido, tipo_contenido, fuentes)
                docx_file = BytesIO()
                doc.save(docx_file)
                docx_file.seek(0)
                st.download_button(
                    label="Descargar contenido como DOCX",
                    data=docx_file,
                    file_name=f"{tipo_contenido.lower().replace(' ', '_')}_{termino.lower().replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        else:
            st.warning("Por favor, selecciona o ingresa un término.")
